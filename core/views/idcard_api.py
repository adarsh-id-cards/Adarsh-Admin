"""
ID Card API Views
Contains: All ID Card Table and ID Card related API endpoints
Including: CRUD, bulk operations, search, status changes, bulk upload
"""
from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import csrf_exempt
import json
import time
import os
from datetime import datetime
from ..models import IDCardGroup, IDCard, IDCardTable


def generate_image_filename(batch_counter, original_ext='.jpg'):
    """
    Generate a unique filename for NEW uploaded images.
    Format: {HHMMSSmmmuuuC}.ext (13 digits total)
    - HH = hours (2 digits)
    - MM = minutes (2 digits)
    - SS = seconds (2 digits)
    - mmm = milliseconds (3 digits)
    - uuu = microseconds (3 digits)
    - C = single digit counter (cycles 1-9, 0, 1-9, 0...)
    
    Example: 1432251234561.jpg (13 digits)
    
    Args:
        batch_counter: The sequential number within the current upload batch (starts from 1)
        original_ext: The original file extension
    
    Returns:
        New filename string (13 digits + extension)
    """
    try:
        # Get file extension
        ext = original_ext.lower() if original_ext else '.jpg'
        if not ext.startswith('.'):
            ext = '.' + ext
        
        # Ensure valid image extension
        valid_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
        if ext not in valid_extensions:
            ext = '.jpg'
        
        # Get current time with microseconds
        now = datetime.now()
        
        # Format: HHMMSS (6 digits)
        time_part = now.strftime('%H%M%S')
        
        # Get milliseconds (3 digits) and microseconds (3 digits)
        microseconds = now.microsecond  # 0-999999
        milliseconds = microseconds // 1000  # First 3 digits (0-999)
        micros = microseconds % 1000  # Last 3 digits (0-999)
        
        mmm = str(milliseconds).zfill(3)
        uuu = str(micros).zfill(3)
        
        # Single digit counter (cycles 1-9, then 0, then repeats)
        counter = batch_counter % 10
        if counter == 0:
            counter = 0  # Keep 0 as 0
        
        # Build filename: HHMMSSmmmuuuC (13 digits)
        filename = f"{time_part}{mmm}{uuu}{counter}{ext}"
        
        return filename
    except Exception as e:
        # Fallback: generate a simple unique filename
        import uuid
        return f"img{uuid.uuid4().hex[:10]}{original_ext or '.jpg'}"


def generate_updated_image_filename(existing_path, new_ext=None):
    """
    Generate updated filename for EXISTING images (update/reupload).
    Keeps the ORIGINAL 13-digit timestamp and adds underscore + 6-digit HHMMSS.
    
    Example: 
    - First upload: 1432251234561.jpg (13 digits)
    - First update: 1432251234561_163045.jpg (13 + 1 + 6 = 20 digits)
    - Second update: 1432251234561_171230.jpg (same 13 digits, new 6-digit time)
    
    Args:
        existing_path: The current file path or filename
        new_ext: Optional new extension (if different from original)
    
    Returns:
        New filename string (20 digits + extension)
    """
    try:
        # Extract just the filename from path
        if existing_path and existing_path not in ['NOT_FOUND', '']:
            filename = os.path.basename(existing_path)
        else:
            # No existing file - generate fresh name
            return generate_image_filename(1, new_ext or '.jpg')
        
        # Get base name and extension
        base_name, current_ext = os.path.splitext(filename)
        
        # Use new extension if provided, otherwise keep current
        ext = new_ext.lower() if new_ext else current_ext.lower()
        if not ext.startswith('.'):
            ext = '.' + ext
        
        # Ensure valid image extension
        valid_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
        if ext not in valid_extensions:
            ext = '.jpg'
        
        # Extract the ORIGINAL 13-digit timestamp from the filename
        # Could be: 1432251234561 (13 digits) or 1432251234561_163045 (20 digits)
        # We need to keep only the first 13 digits
        
        # Remove any underscore suffix first
        if '_' in base_name:
            original_13 = base_name.split('_')[0]
        else:
            original_13 = base_name
        
        # Validate it's 13 digits
        if len(original_13) == 13 and original_13.isdigit():
            # Valid original timestamp
            pass
        else:
            # Filename doesn't follow expected pattern - generate fresh
            return generate_image_filename(1, ext)
        
        # Generate new 6-digit time suffix: HHMMSS
        time_suffix = datetime.now().strftime('%H%M%S')
        
        # Create new filename: original_13_digits + _ + 6_digit_time + ext
        new_filename = f"{original_13}_{time_suffix}{ext}"
        
        return new_filename
    except Exception as e:
        # Fallback: generate a simple unique filename
        import uuid
        return f"updated_{uuid.uuid4().hex[:12]}{new_ext or '.jpg'}"


def get_client_image_folder(client):
    """
    Get the folder path for storing client images.
    Uses client's folder code: 5 chars from name + 5 unique chars.
    Creates the folder if it doesn't exist.
    
    Returns: folder path relative to MEDIA_ROOT like 'adarshimg/{ABCDE12345}/'
    """
    # Ensure folder code is generated
    if not client.image_folder_code:
        client.generate_folder_code()
        client.save(update_fields=['image_folder_code', 'image_folder_suffix'])
    
    folder_path = f"adarshimg/{client.image_folder_code}"
    
    # Ensure folder exists
    from django.conf import settings
    full_path = os.path.join(settings.MEDIA_ROOT, folder_path)
    os.makedirs(full_path, exist_ok=True)
    
    return folder_path


# Counter for unique filenames within same millisecond (legacy - kept for backward compatibility)
_image_upload_counter = 0


def safe_save_image(storage, file_path, file_content, fallback_name=None):
    """
    Safely save an image with fallback handling.
    Returns (saved_path, renamed_filename, success)
    """
    try:
        saved_path = storage.save(file_path, file_content)
        renamed_filename = os.path.basename(file_path)
        return saved_path, renamed_filename, True
    except Exception as e:
        # Try with fallback name if provided
        if fallback_name:
            try:
                fallback_path = os.path.dirname(file_path) + '/' + fallback_name
                saved_path = storage.save(fallback_path, file_content)
                return saved_path, fallback_name, True
            except:
                pass
        return None, None, False


def validate_image_bytes(image_bytes):
    """
    Validate that image bytes represent a valid image.
    Returns (is_valid, error_message)
    """
    try:
        if not image_bytes or len(image_bytes) < 100:
            return False, "Image data is empty or too small"
        
        from PIL import Image
        from io import BytesIO
        
        # Try to open and verify the image
        img = Image.open(BytesIO(image_bytes))
        img.verify()
        
        # Re-open to check it's actually readable
        img = Image.open(BytesIO(image_bytes))
        img.load()  # Force load to catch truncated images
        
        return True, None
    except Exception as e:
        return False, str(e)


# ==================== ID CARD TABLE API ENDPOINTS ====================

@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_table_create(request, group_id):
    """API endpoint to create a new ID Card Table"""
    try:
        group = get_object_or_404(IDCardGroup, id=group_id)
        data = json.loads(request.body)
        
        name = data.get('name', '').strip().upper()  # Convert to uppercase
        if not name:
            return JsonResponse({'success': False, 'message': 'Table name is required!'}, status=400)
        
        fields = data.get('fields', [])
        if len(fields) > 20:
            return JsonResponse({'success': False, 'message': 'Maximum 20 fields allowed!'}, status=400)
        
        # Validate fields structure
        validated_fields = []
        for idx, field in enumerate(fields):
            field_name = field.get('name', '').strip().upper()  # Convert to uppercase
            field_type = field.get('type', 'text')
            if not field_name:
                return JsonResponse({'success': False, 'message': f'Field {idx+1} name is required!'}, status=400)
            
            if field_type not in ['text', 'number', 'date', 'email', 'image', 'textarea']:
                field_type = 'text'
            
            validated_fields.append({
                'name': field_name,
                'type': field_type,
                'order': idx
            })
        
        table = IDCardTable.objects.create(
            group=group,
            name=name,
            fields=validated_fields,
            is_active=True
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Table created successfully!',
            'table': {
                'id': table.id,
                'name': table.name,
                'fields': table.fields,
                'is_active': table.is_active,
                'created_at': table.created_at.strftime('%d-%b-%Y %I:%M %p'),
                'updated_at': table.updated_at.strftime('%d-%b-%Y %I:%M %p'),
            }
        })
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
def api_idcard_table_get(request, table_id):
    """API endpoint to get a single ID Card Table"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        
        return JsonResponse({
            'success': True,
            'table': {
                'id': table.id,
                'name': table.name,
                'fields': table.fields,
                'is_active': table.is_active,
                'created_at': table.created_at.strftime('%d-%b-%Y %I:%M %p'),
                'updated_at': table.updated_at.strftime('%d-%b-%Y %I:%M %p'),
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST", "PUT"])
def api_idcard_table_update(request, table_id):
    """API endpoint to update an ID Card Table"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        data = json.loads(request.body)
        
        name = data.get('name', '').strip().upper()  # Convert to uppercase
        if not name:
            return JsonResponse({'success': False, 'message': 'Table name is required!'}, status=400)
        
        fields = data.get('fields', [])
        if len(fields) > 20:
            return JsonResponse({'success': False, 'message': 'Maximum 20 fields allowed!'}, status=400)
        
        # Validate fields structure
        validated_fields = []
        for idx, field in enumerate(fields):
            field_name = field.get('name', '').strip().upper()  # Convert to uppercase
            field_type = field.get('type', 'text')
            if not field_name:
                return JsonResponse({'success': False, 'message': f'Field {idx+1} name is required!'}, status=400)
            
            if field_type not in ['text', 'number', 'date', 'email', 'image', 'textarea']:
                field_type = 'text'
            
            validated_fields.append({
                'name': field_name,
                'type': field_type,
                'order': idx
            })
        
        table.name = name
        table.fields = validated_fields
        table.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Table updated successfully!',
            'table': {
                'id': table.id,
                'name': table.name,
                'fields': table.fields,
                'is_active': table.is_active,
                'created_at': table.created_at.strftime('%d-%b-%Y %I:%M %p'),
                'updated_at': table.updated_at.strftime('%d-%b-%Y %I:%M %p'),
            }
        })
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["DELETE", "POST"])
def api_idcard_table_delete(request, table_id):
    """API endpoint to delete an ID Card Table"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        table_name = table.name
        table.delete()
        
        return JsonResponse({
            'success': True,
            'message': f'Table "{table_name}" deleted successfully!'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_table_toggle_status(request, table_id):
    """API endpoint to toggle ID Card Table active/inactive status"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        
        if table.is_active:
            table.is_active = False
            status = 'inactive'
            status_display = 'Inactive'
        else:
            table.is_active = True
            status = 'active'
            status_display = 'Active'
        
        table.save()
        
        return JsonResponse({
            'success': True,
            'message': f'Table status changed to {status_display}!',
            'status': status,
            'status_display': status_display
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
def api_idcard_table_list(request, group_id):
    """API endpoint to list all ID Card Tables for a group"""
    try:
        group = get_object_or_404(IDCardGroup, id=group_id)
        tables = IDCardTable.objects.filter(group=group)
        
        table_list = []
        for table in tables:
            table_list.append({
                'id': table.id,
                'name': table.name,
                'fields': table.fields,
                'field_count': len(table.fields),
                'is_active': table.is_active,
                'created_at': table.created_at.strftime('%d-%b-%Y %I:%M %p'),
                'updated_at': table.updated_at.strftime('%d-%b-%Y %I:%M %p'),
            })
        
        return JsonResponse({
            'success': True,
            'tables': table_list
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


# ==================== ID CARD API ENDPOINTS ====================

@csrf_exempt
@require_http_methods(["GET"])
def api_idcard_list(request, table_id):
    """API endpoint to list ID Cards for a table with pagination support for lazy loading"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        status_filter = request.GET.get('status', None)
        
        # Pagination parameters
        offset = int(request.GET.get('offset', 0))
        limit = int(request.GET.get('limit', 100))
        
        # Base queryset ordered by ID
        cards_query = IDCard.objects.filter(table=table).order_by('id')
        if status_filter and status_filter in ['pending', 'verified', 'pool', 'approved', 'download', 'reprint']:
            cards_query = cards_query.filter(status=status_filter)
        
        # Get total count before slicing
        total_count = cards_query.count()
        
        # Apply pagination
        cards = cards_query[offset:offset + limit]
        
        card_list = []
        for idx, card in enumerate(cards):
            # Build ordered fields based on table structure
            ordered_fields = []
            for field in table.fields:
                field_name = field['name']
                field_type = field['type']
                field_value = card.field_data.get(field_name, '')
                ordered_fields.append({
                    'name': field_name,
                    'type': field_type,
                    'value': field_value,
                })
            
            card_list.append({
                'id': card.id,
                'sr_no': offset + idx + 1,
                'field_data': card.field_data,
                'ordered_fields': ordered_fields,
                'photo': card.photo.url if card.photo else None,
                'status': card.status,
                'status_display': card.get_status_display(),
                'created_at': card.created_at.strftime('%d-%b-%Y %I:%M %p'),
                'updated_at': card.updated_at.strftime('%d-%b-%Y %I:%M %p'),
            })
        
        # Get status counts
        status_counts = {
            'pending': IDCard.objects.filter(table=table, status='pending').count(),
            'verified': IDCard.objects.filter(table=table, status='verified').count(),
            'pool': IDCard.objects.filter(table=table, status='pool').count(),
            'approved': IDCard.objects.filter(table=table, status='approved').count(),
            'download': IDCard.objects.filter(table=table, status='download').count(),
            'reprint': IDCard.objects.filter(table=table, status='reprint').count(),
            'total': IDCard.objects.filter(table=table).count(),
        }
        
        return JsonResponse({
            'success': True,
            'cards': card_list,
            'total_count': total_count,
            'offset': offset,
            'limit': limit,
            'has_more': offset + limit < total_count,
            'status_counts': status_counts,
            'table': {
                'id': table.id,
                'name': table.name,
                'fields': table.fields,
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
def api_idcard_all_ids(request, table_id):
    """API endpoint to get all card IDs for a table (for Select All functionality)"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        status_filter = request.GET.get('status', None)
        
        # Base queryset
        cards_query = IDCard.objects.filter(table=table)
        if status_filter and status_filter in ['pending', 'verified', 'pool', 'approved', 'download', 'reprint']:
            cards_query = cards_query.filter(status=status_filter)
        
        # Get only IDs (efficient query)
        card_ids = list(cards_query.values_list('id', flat=True))
        
        return JsonResponse({
            'success': True,
            'card_ids': card_ids,
            'total_count': len(card_ids)
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_create(request, table_id):
    """API endpoint to create a new ID Card with file upload support"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        
        # Get client for folder path
        client = table.group.client
        client_image_folder = get_client_image_folder(client)
        
        # Helper function to convert string values to uppercase
        def uppercase_field_data(data):
            result = {}
            for key, value in data.items():
                if isinstance(value, str):
                    result[key] = value.upper()
                else:
                    result[key] = value
            return result
        
        # Handle both JSON and FormData
        if request.content_type and 'multipart/form-data' in request.content_type:
            # FormData submission (with files)
            field_data_str = request.POST.get('field_data', '{}')
            field_data = json.loads(field_data_str)
            field_data = uppercase_field_data(field_data)
            
            # Handle image fields from table configuration
            image_counter = 0
            for field in table.fields:
                if field['type'] == 'image':
                    field_name = field['name']
                    file_key = f"image_{field_name}"
                    if file_key in request.FILES:
                        try:
                            # Save the image with new filename (14-digit timestamp + counter) for NEW cards
                            uploaded_file = request.FILES[file_key]
                            from django.core.files.storage import default_storage
                            
                            # Get file extension
                            original_ext = os.path.splitext(uploaded_file.name)[1].lower() or '.jpg'
                            
                            # Generate new filename with 14-digit timestamp + counter (for NEW cards)
                            image_counter += 1
                            new_filename = generate_image_filename(image_counter, original_ext)
                            file_path = f"{client_image_folder}/{new_filename}"
                            
                            # Try to save the image
                            saved_path, renamed, success = safe_save_image(
                                default_storage, file_path, uploaded_file, 
                                fallback_name=f"fallback_{int(time.time())}.jpg"
                            )
                            
                            if success and saved_path:
                                field_data[field_name] = saved_path
                            else:
                                # Log the error but continue - don't break the whole operation
                                print(f"Warning: Could not save image for field {field_name}")
                        except Exception as img_err:
                            # Log error but continue with other fields
                            print(f"Error processing image for field {field_name}: {img_err}")
            
            # Create the card
            card = IDCard.objects.create(
                table=table,
                field_data=field_data,
                status='pending'
            )
            
            # Handle main photo
            if 'photo' in request.FILES:
                card.photo = request.FILES['photo']
                card.save()
        else:
            # JSON submission (no files)
            data = json.loads(request.body)
            field_data = data.get('field_data', {})
            field_data = uppercase_field_data(field_data)
            
            card = IDCard.objects.create(
                table=table,
                field_data=field_data,
                status='pending'
            )
        
        return JsonResponse({
            'success': True,
            'message': 'ID Card created successfully!',
            'card': {
                'id': card.id,
                'field_data': card.field_data,
                'photo': card.photo.url if card.photo else None,
                'status': card.status,
                'created_at': card.created_at.strftime('%d-%b-%Y %I:%M %p'),
            }
        })
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
def api_idcard_get(request, card_id):
    """API endpoint to get a single ID Card"""
    try:
        card = get_object_or_404(IDCard, id=card_id)
        
        return JsonResponse({
            'success': True,
            'card': {
                'id': card.id,
                'table_id': card.table.id,
                'table_name': card.table.name,
                'field_data': card.field_data,
                'photo': card.photo.url if card.photo else None,
                'status': card.status,
                'status_display': card.get_status_display(),
                'created_at': card.created_at.strftime('%d-%b-%Y %I:%M %p'),
                'updated_at': card.updated_at.strftime('%d-%b-%Y %I:%M %p'),
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST", "PUT"])
def api_idcard_update(request, card_id):
    """API endpoint to update an ID Card with file upload support"""
    try:
        card = get_object_or_404(IDCard, id=card_id)
        table = card.table
        
        # Get client for folder path
        client = table.group.client
        client_image_folder = get_client_image_folder(client)
        
        # Helper function to convert string values to uppercase
        def uppercase_field_data(data):
            result = {}
            for key, value in data.items():
                if isinstance(value, str):
                    result[key] = value.upper()
                else:
                    result[key] = value
            return result
        
        # Handle both JSON and FormData
        if request.content_type and 'multipart/form-data' in request.content_type:
            # FormData submission (with files)
            field_data_str = request.POST.get('field_data', '{}')
            new_field_data = json.loads(field_data_str)
            new_field_data = uppercase_field_data(new_field_data)
            
            # Merge with existing field_data to preserve existing image paths
            existing_field_data = card.field_data or {}
            existing_field_data.update(new_field_data)
            
            # Handle image fields from table configuration
            image_counter = 0
            for field in table.fields:
                if field['type'] == 'image':
                    field_name = field['name']
                    file_key = f"image_{field_name}"
                    if file_key in request.FILES:
                        try:
                            # Save the image with new filename
                            uploaded_file = request.FILES[file_key]
                            from django.core.files.storage import default_storage
                            
                            # Get file extension from uploaded file
                            original_ext = os.path.splitext(uploaded_file.name)[1].lower() or '.jpg'
                            
                            # Check if there's an existing image for this field
                            existing_image_path = existing_field_data.get(field_name, '')
                            
                            # Generate filename based on whether image already exists
                            if existing_image_path and existing_image_path != 'NOT_FOUND' and existing_image_path.strip():
                                # UPDATE: Keep original 13-digit timestamp, add new 6-digit HHMMSS (20 digits total)
                                new_filename = generate_updated_image_filename(existing_image_path, original_ext)
                                
                                # Delete old image
                                try:
                                    if default_storage.exists(existing_image_path):
                                        default_storage.delete(existing_image_path)
                                        print(f"Deleted old image: {existing_image_path}")
                                except Exception as del_err:
                                    print(f"Warning: Could not delete old image {existing_image_path}: {del_err}")
                            else:
                                # FIRST UPLOAD: Generate fresh 13-digit filename
                                image_counter += 1
                                new_filename = generate_image_filename(image_counter, original_ext)
                            
                            file_path = f"{client_image_folder}/{new_filename}"
                            
                            # Try to save the image
                            saved_path, renamed, success = safe_save_image(
                                default_storage, file_path, uploaded_file,
                                fallback_name=f"fallback_{int(time.time())}.jpg"
                            )
                            
                            if success and saved_path:
                                existing_field_data[field_name] = saved_path
                            else:
                                # Log the error but continue
                                print(f"Warning: Could not save image for field {field_name}")
                        except Exception as img_err:
                            # Log error but continue with other fields
                            print(f"Error processing image for field {field_name}: {img_err}")
            
            card.field_data = existing_field_data
            
            # Handle main photo - save to field_data["Photo"] not card.photo
            if 'photo' in request.FILES:
                try:
                    uploaded_file = request.FILES['photo']
                    from django.core.files.storage import default_storage
                    
                    # Get file extension from uploaded file
                    original_ext = os.path.splitext(uploaded_file.name)[1].lower() or '.jpg'
                    
                    # Check if there's an existing Photo in field_data
                    existing_photo_path = existing_field_data.get('Photo', '')
                    
                    # Generate filename based on whether photo already exists
                    if existing_photo_path and existing_photo_path != 'NOT_FOUND' and existing_photo_path.strip():
                        # UPDATE: Keep original 13-digit timestamp, add new 6-digit HHMMSS (20 digits total)
                        new_filename = generate_updated_image_filename(existing_photo_path, original_ext)
                        
                        # Delete old photo
                        try:
                            if default_storage.exists(existing_photo_path):
                                default_storage.delete(existing_photo_path)
                                print(f"Deleted old photo: {existing_photo_path}")
                        except Exception as del_err:
                            print(f"Warning: Could not delete old photo {existing_photo_path}: {del_err}")
                    else:
                        # FIRST UPLOAD: Generate fresh 13-digit filename
                        new_filename = generate_image_filename(9, original_ext)  # Use 9 as counter for main photo
                    
                    file_path = f"{client_image_folder}/{new_filename}"
                    
                    # Save the image
                    saved_path, renamed, success = safe_save_image(
                        default_storage, file_path, uploaded_file,
                        fallback_name=f"photo_{int(time.time())}.jpg"
                    )
                    
                    if success and saved_path:
                        existing_field_data['Photo'] = saved_path
                        card.field_data = existing_field_data
                        print(f"Photo updated to: {saved_path}")
                    else:
                        print(f"Warning: Could not save main photo")
                except Exception as photo_err:
                    print(f"Error processing main photo: {photo_err}")
            
            card.save()
        else:
            # JSON submission (no files)
            data = json.loads(request.body)
            
            if 'field_data' in data:
                # Merge with existing field_data to preserve image paths and other fields
                existing_field_data = card.field_data or {}
                new_field_data = uppercase_field_data(data['field_data'])
                existing_field_data.update(new_field_data)
                card.field_data = existing_field_data
            if 'status' in data and data['status'] in ['pending', 'verified', 'pool', 'approved', 'download', 'reprint']:
                card.status = data['status']
            
            card.save()
        
        return JsonResponse({
            'success': True,
            'message': 'ID Card updated successfully!',
            'card': {
                'id': card.id,
                'field_data': card.field_data,
                'photo': card.photo.url if card.photo else None,
                'status': card.status,
                'status_display': card.get_status_display(),
            }
        })
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["DELETE", "POST"])
def api_idcard_delete(request, card_id):
    """API endpoint to delete an ID Card"""
    try:
        card = get_object_or_404(IDCard, id=card_id)
        card.delete()
        
        return JsonResponse({
            'success': True,
            'message': 'ID Card deleted successfully!'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_update_field(request, card_id):
    """API endpoint to update a single field on an ID Card (for inline editing)"""
    try:
        card = get_object_or_404(IDCard, id=card_id)
        data = json.loads(request.body)
        
        field = data.get('field')
        value = data.get('value', '')
        
        if not field:
            return JsonResponse({'success': False, 'message': 'Field name is required!'}, status=400)
        
        # Update the field_data
        field_data = card.field_data or {}
        
        # Convert to uppercase if it's a string
        if isinstance(value, str):
            field_data[field] = value.upper()
        else:
            field_data[field] = value
        
        card.field_data = field_data
        card.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Field updated successfully!',
            'field': field,
            'value': field_data[field]
        })
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_change_status(request, card_id):
    """API endpoint to change an ID Card's status"""
    try:
        card = get_object_or_404(IDCard, id=card_id)
        data = json.loads(request.body)
        
        new_status = data.get('status')
        if new_status not in ['pending', 'verified', 'pool', 'approved', 'download', 'reprint']:
            return JsonResponse({'success': False, 'message': 'Invalid status!'}, status=400)
        
        card.status = new_status
        card.save()
        
        return JsonResponse({
            'success': True,
            'message': f'Card status changed to {card.get_status_display()}!',
            'status': card.status,
            'status_display': card.get_status_display()
        })
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_bulk_status(request, table_id):
    """API endpoint to change status of multiple ID Cards"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        data = json.loads(request.body)
        
        card_ids = data.get('card_ids', [])
        new_status = data.get('status')
        
        if new_status not in ['pending', 'verified', 'pool', 'approved', 'download', 'reprint']:
            return JsonResponse({'success': False, 'message': 'Invalid status!'}, status=400)
        
        updated_count = IDCard.objects.filter(table=table, id__in=card_ids).update(status=new_status)
        
        return JsonResponse({
            'success': True,
            'message': f'{updated_count} cards updated to {new_status}!',
            'updated_count': updated_count
        })
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_bulk_delete(request, table_id):
    """API endpoint to delete multiple ID Cards"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        data = json.loads(request.body)
        
        card_ids = data.get('card_ids', [])
        delete_all = data.get('delete_all', False)
        
        if delete_all:
            deleted_count, _ = IDCard.objects.filter(table=table).delete()
        else:
            deleted_count, _ = IDCard.objects.filter(table=table, id__in=card_ids).delete()
        
        return JsonResponse({
            'success': True,
            'message': f'{deleted_count} cards deleted successfully!',
            'deleted_count': deleted_count
        })
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
def api_idcard_search(request, table_id):
    """API endpoint to search ID Cards across all statuses"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        query = request.GET.get('q', '').strip().upper()
        
        if not query or len(query) < 2:
            return JsonResponse({
                'success': True,
                'results': [],
                'message': 'Please enter at least 2 characters to search'
            })
        
        # Get all cards for this table
        cards = IDCard.objects.filter(table=table)
        
        results = []
        for card in cards:
            # Search in field_data
            field_data = card.field_data or {}
            match_found = False
            matched_field = ''
            matched_value = ''
            
            for field_name, field_value in field_data.items():
                if field_value and query in str(field_value).upper():
                    match_found = True
                    matched_field = field_name
                    matched_value = str(field_value)
                    break
            
            if match_found:
                # Get the first text field as display name
                display_name = ''
                for field in table.fields:
                    if field.get('type') in ['text', 'textarea'] and field.get('name') in field_data:
                        display_name = field_data.get(field.get('name'), '')
                        break
                
                results.append({
                    'id': card.id,
                    'display_name': display_name or f'Card #{card.id}',
                    'status': card.status,
                    'status_display': card.get_status_display(),
                    'matched_field': matched_field,
                    'matched_value': matched_value,
                    'photo': card.photo.url if card.photo else None,
                    'field_data': card.field_data,
                })
        
        return JsonResponse({
            'success': True,
            'results': results,
            'count': len(results),
            'query': query
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
def api_table_status_counts(request, table_id):
    """API endpoint to get status counts for a table"""
    try:
        table = get_object_or_404(IDCardTable, id=table_id)
        
        status_counts = {
            'pending': IDCard.objects.filter(table=table, status='pending').count(),
            'verified': IDCard.objects.filter(table=table, status='verified').count(),
            'pool': IDCard.objects.filter(table=table, status='pool').count(),
            'approved': IDCard.objects.filter(table=table, status='approved').count(),
            'download': IDCard.objects.filter(table=table, status='download').count(),
            'reprint': IDCard.objects.filter(table=table, status='reprint').count(),
            'total': IDCard.objects.filter(table=table).count(),
        }
        
        return JsonResponse({
            'success': True,
            'status_counts': status_counts
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_bulk_upload(request, table_id):
    """API endpoint to bulk upload ID Cards from XLSX/CSV file with fuzzy matching and optional ZIP photo upload"""
    try:
        import openpyxl
        from io import BytesIO
        import re
        import zipfile
        import os
        from django.core.files.storage import default_storage
        from django.core.files.base import ContentFile
        
        table = get_object_or_404(IDCardTable, id=table_id)
        
        if 'file' not in request.FILES:
            return JsonResponse({'success': False, 'message': 'No file uploaded!'}, status=400)
        
        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name.lower()
        file_size = uploaded_file.size
        
        # Check if ZIP file with photos is also uploaded
        photos_zip_file = request.FILES.get('photos_zip', None)
        zip_photos = {}  # Dictionary to map filename (without extension) to image bytes - CASE SENSITIVE
        
        if photos_zip_file:
            try:
                zip_content = photos_zip_file.read()
                with zipfile.ZipFile(BytesIO(zip_content), 'r') as zf:
                    for zip_info in zf.infolist():
                        # Skip directories
                        if zip_info.is_dir():
                            continue
                        
                        file_in_zip = zip_info.filename
                        # Get filename without path and extension
                        base_name = os.path.basename(file_in_zip)
                        # CASE SENSITIVE - keep original case, no extension
                        name_without_ext = os.path.splitext(base_name)[0]
                        ext = os.path.splitext(base_name)[1].lower()
                        
                        # Only process image files
                        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
                            try:
                                image_bytes = zf.read(zip_info.filename)
                                
                                # Validate image before storing
                                is_valid, error_msg = validate_image_bytes(image_bytes)
                                if is_valid:
                                    zip_photos[name_without_ext] = {
                                        'bytes': image_bytes,
                                        'ext': ext
                                    }
                                else:
                                    print(f"Invalid image skipped in ZIP: {base_name} - {error_msg}")
                            except Exception as img_read_err:
                                print(f"Error reading image from ZIP: {base_name} - {img_read_err}")
                                continue
            except Exception as zip_error:
                # Continue without photos - don't fail the entire upload
                pass
        
        # Get client for image folder path
        client = table.group.client
        client_image_folder = get_client_image_folder(client)
        
        # Get all table fields (text fields for matching, image fields to include with empty values)
        all_table_fields = table.fields
        table_fields = [f['name'] for f in all_table_fields if f.get('type') != 'image']
        image_fields = [f['name'] for f in all_table_fields if f.get('type') == 'image']
        
        # Also check if there's a PHOTO field (even if not marked as image type)
        # This allows matching ZIP photos to a PHOTO column
        photo_field_name = None
        for f in all_table_fields:
            if f['name'].upper() == 'PHOTO':
                photo_field_name = f['name']
                if photo_field_name not in image_fields:
                    image_fields.append(photo_field_name)
                # Remove from table_fields if present (we'll handle it as image)
                if photo_field_name in table_fields:
                    table_fields.remove(photo_field_name)
                break
        
        # Helper function to normalize field names for comparison
        def normalize_name(name):
            # Remove spaces, underscores, hyphens, dots and lowercase
            return re.sub(r'[^a-z0-9]', '', name.lower())
        
        # Levenshtein distance for fuzzy matching
        def levenshtein_distance(s1, s2):
            if len(s1) < len(s2):
                return levenshtein_distance(s2, s1)
            if len(s2) == 0:
                return len(s1)
            
            previous_row = range(len(s2) + 1)
            for i, c1 in enumerate(s1):
                current_row = [i + 1]
                for j, c2 in enumerate(s2):
                    insertions = previous_row[j + 1] + 1
                    deletions = current_row[j] + 1
                    substitutions = previous_row[j] + (c1 != c2)
                    current_row.append(min(insertions, deletions, substitutions))
                previous_row = current_row
            return previous_row[-1]
        
        # Find best match for a header using fuzzy matching
        def find_best_match(header, available_fields):
            normalized_header = normalize_name(header)
            
            # First try exact match
            for field in available_fields:
                if normalize_name(field) == normalized_header:
                    return field
            
            # Then try fuzzy match
            best_match = None
            best_distance = float('inf')
            
            for field in available_fields:
                normalized_field = normalize_name(field)
                distance = levenshtein_distance(normalized_header, normalized_field)
                
                # Allow up to 2 char differences, but for short strings allow only 1
                max_distance = 1 if len(normalized_field) < 5 else 2
                
                if distance <= max_distance and distance < best_distance:
                    best_distance = distance
                    best_match = field
            
            return best_match
        
        cards_created = 0
        total_photos_matched = 0
        errors = []
        matched_field_names = []
        
        if file_name.endswith('.xlsx') or file_name.endswith('.xls'):
            # Process Excel file
            try:
                file_content = uploaded_file.read()
                
                # Check file magic bytes to detect actual format
                if len(file_content) < 4:
                    return JsonResponse({
                        'success': False, 
                        'message': 'File is too small or empty.'
                    }, status=400)
                
                magic_bytes = file_content[:4]
                is_zip = magic_bytes[:2] == b'PK'
                is_old_xls = (magic_bytes[0] == 0xD0 and magic_bytes[1] == 0xCF)
                
                headers = []
                rows_data = []
                
                if is_zip or file_name.endswith('.xlsx'):
                    # New .xlsx format - use openpyxl
                    try:
                        wb = openpyxl.load_workbook(BytesIO(file_content))
                        ws = wb.active
                        
                        # Get header row (first row)
                        for cell in ws[1]:
                            if cell.value:
                                headers.append(str(cell.value).strip())
                        
                        # Get data rows
                        for row in ws.iter_rows(min_row=2, values_only=True):
                            rows_data.append(row)
                    except Exception as xlsx_error:
                        # Maybe it's actually an old xls file with wrong extension
                        if not is_zip:
                            is_old_xls = True
                        else:
                            raise xlsx_error
                
                if is_old_xls or (file_name.endswith('.xls') and not file_name.endswith('.xlsx') and not headers):
                    # Old .xls format - use xlrd
                    try:
                        import xlrd
                        wb = xlrd.open_workbook(file_contents=file_content)
                        ws = wb.sheet_by_index(0)
                        
                        # Get header row (first row)
                        headers = []
                        for col_idx in range(ws.ncols):
                            cell_value = ws.cell_value(0, col_idx)
                            if cell_value:
                                headers.append(str(cell_value).strip())
                        
                        # Get data rows
                        rows_data = []
                        for row_idx in range(1, ws.nrows):
                            row = []
                            for col_idx in range(ws.ncols):
                                row.append(ws.cell_value(row_idx, col_idx))
                            rows_data.append(tuple(row))
                    except ImportError:
                        return JsonResponse({
                            'success': False, 
                            'message': 'xlrd library not installed. Please install it to support .xls files.'
                        }, status=400)
                    except Exception as xls_error:
                        return JsonResponse({
                            'success': False, 
                            'message': f'Error reading .xls file: {str(xls_error)}'
                        }, status=400)
                
                if not headers:
                    return JsonResponse({
                        'success': False, 
                        'message': 'Could not read headers from Excel file. Please check the file format.'
                    }, status=400)
                    
            except Exception as excel_error:
                return JsonResponse({
                    'success': False, 
                    'message': f'Error reading Excel file: {str(excel_error)}'
                }, status=400)
            
            # Map headers to table fields using fuzzy matching
            # Skip any headers that match image field names (those are for ZIP matching only)
            header_to_field = {}
            available_fields = table_fields.copy()
            image_field_names_upper = [f.upper() for f in image_fields]
            
            # Track which column indices contain image reference values (like PHOTO column)
            image_ref_columns = {}
            photo_column_idx = None  # Track the generic PHOTO column for any image field
            
            for idx, header in enumerate(headers):
                header_upper = header.upper() if header else ''
                
                # Check if this is a PHOTO column (generic image reference)
                if header_upper == 'PHOTO':
                    photo_column_idx = idx
                    # Assign PHOTO column to first image field if not already assigned
                    if image_fields and image_fields[0] not in image_ref_columns:
                        image_ref_columns[image_fields[0]] = idx
                    continue  # Skip text field matching for PHOTO column
                
                # Check if this column matches any image field name exactly
                is_image_ref = False
                for img_field in image_fields:
                    if header_upper == img_field.upper():
                        image_ref_columns[img_field] = idx
                        is_image_ref = True
                        break
                
                # Skip this column for text field matching if it's an image reference column
                if is_image_ref:
                    continue
                    
                match = find_best_match(header, available_fields)
                if match:
                    header_to_field[idx] = match
                    available_fields.remove(match)  # Don't match same field twice
                    matched_field_names.append(match)
            
            if not header_to_field:
                return JsonResponse({
                    'success': False, 
                    'message': f'No matching columns found! Expected columns: {", ".join(table_fields)}'
                }, status=400)
            
            # Process data rows using rows_data collected earlier
            for row_num, row in enumerate(rows_data, start=2):
                try:
                    # Skip empty rows
                    if all(cell is None or str(cell).strip() == '' for cell in row):
                        continue
                    
                    field_data = {}
                    for col_idx, field_name in header_to_field.items():
                        if col_idx < len(row):
                            value = row[col_idx]
                            if value is not None:
                                # Convert to string, handle dates and numbers
                                if hasattr(value, 'strftime'):
                                    # Already a datetime object
                                    value = value.strftime('%d-%m-%Y')
                                elif isinstance(value, float):
                                    # Check if it's an Excel date serial number (typically between 1 and 60000)
                                    # Excel dates start from 1900-01-01 (serial 1)
                                    if 1 < value < 60000 and ('date' in field_name.lower() or 'dob' in field_name.lower() or 'birth' in field_name.lower()):
                                        # Convert Excel serial date to actual date
                                        from datetime import datetime, timedelta
                                        # Excel's epoch is December 30, 1899
                                        excel_epoch = datetime(1899, 12, 30)
                                        actual_date = excel_epoch + timedelta(days=int(value))
                                        value = actual_date.strftime('%d-%m-%Y')
                                    elif value == int(value):
                                        # It's actually an integer (no decimal part)
                                        value = str(int(value)).upper()
                                    else:
                                        value = str(value).upper()
                                elif isinstance(value, int):
                                    # Check if it might be an Excel date serial for integer values
                                    if 1 < value < 60000 and ('date' in field_name.lower() or 'dob' in field_name.lower() or 'birth' in field_name.lower()):
                                        from datetime import datetime, timedelta
                                        excel_epoch = datetime(1899, 12, 30)
                                        actual_date = excel_epoch + timedelta(days=value)
                                        value = actual_date.strftime('%d-%m-%Y')
                                    else:
                                        value = str(value).upper()
                                else:
                                    value = str(value).strip().upper()  # Convert to uppercase
                                field_data[field_name] = value
                            else:
                                field_data[field_name] = ''
                        else:
                            field_data[field_name] = ''
                    
                    # Process image fields - try to match with ZIP photos
                    photos_matched = 0
                    
                    for img_field in image_fields:
                        # Get the photo reference value from the tracked column
                        photo_column_value = None
                        
                        # Use the tracked image reference column if available
                        if img_field in image_ref_columns:
                            col_idx = image_ref_columns[img_field]
                            if col_idx < len(row):
                                cell_value = row[col_idx]
                                if cell_value is not None:
                                    # Handle numeric values - convert float 1.0 to "1"
                                    # CASE SENSITIVE - do NOT convert to uppercase
                                    if isinstance(cell_value, float) and cell_value == int(cell_value):
                                        photo_column_value = str(int(cell_value))
                                    elif isinstance(cell_value, int):
                                        photo_column_value = str(cell_value)
                                    else:
                                        photo_column_value = str(cell_value).strip()
                        else:
                            # Fallback: look for column named PHOTO or use photo_column_idx
                            ref_col_idx = photo_column_idx if photo_column_idx is not None else None
                            if ref_col_idx is None:
                                # Search for PHOTO column
                                for col_idx in range(len(headers)):
                                    if headers[col_idx] and headers[col_idx].upper() == 'PHOTO':
                                        ref_col_idx = col_idx
                                        break
                            
                            if ref_col_idx is not None and ref_col_idx < len(row):
                                cell_value = row[ref_col_idx]
                                if cell_value is not None:
                                    # Handle numeric values - CASE SENSITIVE
                                    if isinstance(cell_value, float) and cell_value == int(cell_value):
                                        photo_column_value = str(int(cell_value))
                                    elif isinstance(cell_value, int):
                                        photo_column_value = str(cell_value)
                                    else:
                                        photo_column_value = str(cell_value).strip()
                        
                        # Try to match photo from ZIP - CASE SENSITIVE match
                        if photo_column_value and zip_photos and photo_column_value in zip_photos:
                            try:
                                photo_info = zip_photos[photo_column_value]
                                
                                # Generate new filename with 14-digit timestamp + batch counter
                                cards_created += 1  # Increment before for 1-based counter
                                original_ext = photo_info['ext']
                                new_filename = generate_image_filename(cards_created, original_ext)
                                
                                # Use client's unique folder: id_card_images/{client_uuid}/
                                file_path = f"{client_image_folder}/{new_filename}"
                                
                                # Save the image with error handling
                                saved_path, renamed, success = safe_save_image(
                                    default_storage, file_path, ContentFile(photo_info['bytes']),
                                    fallback_name=f"fallback_{int(time.time())}{original_ext}"
                                )
                                
                                if success and saved_path:
                                    field_data[img_field] = saved_path
                                    photos_matched += 1
                                    total_photos_matched += 1
                                else:
                                    field_data[img_field] = 'NOT_FOUND'  # Save failed
                                cards_created -= 1  # Revert since we incremented early
                            except Exception as photo_error:
                                # Log but don't break the whole process
                                print(f"Error saving photo (XLSX) for {photo_column_value}: {photo_error}")
                                field_data[img_field] = f'PENDING:{photo_column_value}'  # Store reference for later
                        elif photo_column_value:
                            # Value given but image not found in ZIP - store reference for reupload matching
                            field_data[img_field] = f'PENDING:{photo_column_value}'
                        else:
                            # No value given at all - show colorful placeholder
                            field_data[img_field] = ''
                    
                    # Create the card
                    IDCard.objects.create(
                        table=table,
                        field_data=field_data,
                        status='pending'
                    )
                    cards_created += 1
                    
                except Exception as e:
                    errors.append(f'Row {row_num}: {str(e)}')
        
        elif file_name.endswith('.csv'):
            import csv
            from io import StringIO
            
            # Read CSV
            content = uploaded_file.read().decode('utf-8-sig')
            reader = csv.DictReader(StringIO(content))
            
            # Map CSV headers to table fields using fuzzy matching
            csv_headers = reader.fieldnames or []
            header_to_field = {}
            available_fields = table_fields.copy()
            
            # Track image reference columns for CSV
            csv_image_ref_columns = {}
            csv_photo_column = None  # Track the generic PHOTO column
            
            for header in csv_headers:
                header_upper = header.upper() if header else ''
                
                # Check if this is a PHOTO column
                if header_upper == 'PHOTO':
                    csv_photo_column = header
                    # Assign to first image field if not already assigned
                    if image_fields and image_fields[0] not in csv_image_ref_columns:
                        csv_image_ref_columns[image_fields[0]] = header
                    continue
                
                # Check if this column matches any image field name exactly
                is_image_ref = False
                for img_field in image_fields:
                    if header_upper == img_field.upper():
                        csv_image_ref_columns[img_field] = header
                        is_image_ref = True
                        break
                
                # Skip this column for text field matching if it's an image reference column
                if is_image_ref:
                    continue
                    
                match = find_best_match(header.strip(), available_fields)
                if match:
                    header_to_field[header] = match
                    available_fields.remove(match)
                    matched_field_names.append(match)
            
            if not header_to_field:
                return JsonResponse({
                    'success': False, 
                    'message': f'No matching columns found! Expected columns: {", ".join(table_fields)}'
                }, status=400)
            
            for row_num, row in enumerate(reader, start=2):
                try:
                    # Skip empty rows
                    if all(not v or str(v).strip() == '' for v in row.values()):
                        continue
                    
                    field_data = {}
                    for csv_header, field_name in header_to_field.items():
                        value = row.get(csv_header, '')
                        field_data[field_name] = str(value).strip().upper() if value else ''  # Convert to uppercase
                    
                    # Process image fields - try to match with ZIP photos
                    photos_matched = 0
                    
                    for img_field in image_fields:
                        # Get the photo reference value from the tracked column
                        photo_column_value = None
                        
                        # Use tracked image reference column if available
                        if img_field in csv_image_ref_columns:
                            csv_header = csv_image_ref_columns[img_field]
                            cell_value = row.get(csv_header, '')
                            if cell_value and str(cell_value).strip():
                                # CASE SENSITIVE - do NOT convert to uppercase
                                photo_column_value = str(cell_value).strip()
                        else:
                            # Fallback: look for column named PHOTO or matching image field name
                            for csv_header, cell_value in row.items():
                                header_upper = csv_header.upper() if csv_header else ''
                                if header_upper == 'PHOTO' or header_upper == img_field.upper():
                                    if cell_value and str(cell_value).strip():
                                        # CASE SENSITIVE
                                        photo_column_value = str(cell_value).strip()
                                    break
                        
                        # Try to match photo from ZIP - CASE SENSITIVE match
                        if photo_column_value and zip_photos and photo_column_value in zip_photos:
                            try:
                                photo_info = zip_photos[photo_column_value]
                                
                                # Generate new filename with 14-digit timestamp + batch counter
                                cards_created += 1  # Increment before for 1-based counter
                                original_ext = photo_info['ext']
                                new_filename = generate_image_filename(cards_created, original_ext)
                                
                                # Use client's unique folder: id_card_images/{client_uuid}/
                                file_path = f"{client_image_folder}/{new_filename}"
                                
                                # Save the image with error handling
                                saved_path, renamed, success = safe_save_image(
                                    default_storage, file_path, ContentFile(photo_info['bytes']),
                                    fallback_name=f"fallback_{int(time.time())}{original_ext}"
                                )
                                
                                if success and saved_path:
                                    field_data[img_field] = saved_path
                                    photos_matched += 1
                                    total_photos_matched += 1
                                else:
                                    field_data[img_field] = f'PENDING:{photo_column_value}'  # Save failed, store reference
                                cards_created -= 1  # Revert since we incremented early
                            except Exception as photo_error:
                                # Log but don't break the whole process
                                print(f"Error saving photo (CSV) for {photo_column_value}: {photo_error}")
                                field_data[img_field] = f'PENDING:{photo_column_value}'  # Store reference for later
                        elif photo_column_value:
                            # Value given but image not found in ZIP - store reference for reupload matching
                            field_data[img_field] = f'PENDING:{photo_column_value}'
                        else:
                            # No value given at all - show colorful placeholder
                            field_data[img_field] = ''
                    
                    # Create the card
                    IDCard.objects.create(
                        table=table,
                        field_data=field_data,
                        status='pending'
                    )
                    cards_created += 1
                    
                except Exception as e:
                    errors.append(f'Row {row_num}: {str(e)}')
        
        else:
            return JsonResponse({
                'success': False, 
                'message': 'Invalid file format! Please upload .xlsx, .xls, or .csv file.'
            }, status=400)
        
        # Return result
        photo_msg = f" with {total_photos_matched} photos matched" if total_photos_matched > 0 else ""
        result = {
            'success': True,
            'message': f'Successfully created {cards_created} ID cards{photo_msg}!',
            'cards_created': cards_created,
            'photos_matched': total_photos_matched,
            'matched_fields': matched_field_names,
        }
        
        if errors:
            result['errors'] = errors[:10]  # Return first 10 errors only
            result['error_count'] = len(errors)
        
        return JsonResponse(result)
        
    except ImportError:
        return JsonResponse({
            'success': False, 
            'message': 'openpyxl library not installed. Run: pip install openpyxl'
        }, status=500)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_download_images(request, table_id):
    """API endpoint to download images as ZIP for selected cards"""
    try:
        import zipfile
        import os
        from io import BytesIO
        from datetime import datetime
        from django.http import HttpResponse
        from django.core.files.storage import default_storage
        
        table = get_object_or_404(IDCardTable, id=table_id)
        data = json.loads(request.body)
        
        card_ids = data.get('card_ids', [])
        if not card_ids:
            return JsonResponse({'success': False, 'message': 'No cards selected!'}, status=400)
        
        # Get selected cards in database order
        cards = IDCard.objects.filter(table=table, id__in=card_ids).order_by('id')
        if not cards.exists():
            return JsonResponse({'success': False, 'message': 'No cards found!'}, status=400)
        
        # Generate timestamp for filenames (format: YYYYMMDD_HHMMSS)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Get image field names from table config
        image_fields = [f['name'] for f in table.fields if f.get('type') == 'image']
        
        # Also check for PHOTO field
        for f in table.fields:
            if f['name'].upper() == 'PHOTO' and f['name'] not in image_fields:
                image_fields.append(f['name'])
        
        # Create ZIP file in memory
        zip_buffer = BytesIO()
        images_added = 0
        images_skipped = 0
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for card in cards:
                field_data = card.field_data or {}
                
                # Check all image fields
                for img_field in image_fields:
                    img_path = field_data.get(img_field, '')
                    
                    if img_path and img_path != 'NOT_FOUND' and img_path.strip():
                        try:
                            # Get file from storage
                            if default_storage.exists(img_path):
                                with default_storage.open(img_path, 'rb') as img_file:
                                    img_data = img_file.read()
                                    
                                    # Validate image before adding to ZIP
                                    if img_data and len(img_data) >= 100:
                                        # Use the stored filename (already renamed with timestamp format)
                                        download_filename = os.path.basename(img_path)
                                        
                                        # Add to ZIP with the filename
                                        zf.writestr(download_filename, img_data)
                                        images_added += 1
                                    else:
                                        images_skipped += 1
                                        print(f"Skipping empty/small image: {img_path}")
                            else:
                                images_skipped += 1
                        except Exception as e:
                            # Skip this image if error but continue with others
                            images_skipped += 1
                            print(f"Error downloading image {img_path}: {e}")
                            continue
        
        if images_added == 0:
            return JsonResponse({'success': False, 'message': 'No images found for selected cards!'}, status=400)
        
        # Create response with ZIP file
        zip_buffer.seek(0)
        
        # Create ZIP filename with table name and timestamp
        zip_filename = f"{table.name}_{timestamp}_images.zip"
        
        response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{zip_filename}"'
        response['Content-Length'] = len(zip_buffer.getvalue())
        
        return response
        
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_reupload_images(request, table_id):
    """
    API endpoint to reupload images from a ZIP file.
    Matches ZIP filenames (without extension) to the current stored filename on each IDCard.
    On match, saves new image with versioned filename (adds _XXXXXX suffix) in client's UUID folder.
    Uses CASE SENSITIVE matching.
    """
    try:
        import zipfile
        from io import BytesIO
        from django.core.files.storage import default_storage
        from django.core.files.base import ContentFile
        
        table = get_object_or_404(IDCardTable, id=table_id)
        
        # Get client for folder path
        client = table.group.client
        client_image_folder = get_client_image_folder(client)
        
        if 'zip_file' not in request.FILES:
            return JsonResponse({'success': False, 'message': 'No ZIP file uploaded!'}, status=400)
        
        # Parse card_ids from request
        card_ids = []
        if 'card_ids' in request.POST:
            try:
                card_ids = json.loads(request.POST.get('card_ids', '[]'))
            except:
                card_ids = []
        
        # Get cards to process in database order
        if card_ids:
            cards = IDCard.objects.filter(table=table, id__in=card_ids).order_by('id')
        else:
            cards = IDCard.objects.filter(table=table).order_by('id')
        
        # Get all cards - match by current stored filename
        cards_to_process = list(cards)
        
        if not cards_to_process:
            return JsonResponse({
                'success': False, 
                'message': 'No cards found to process!'
            }, status=400)
        
        # Get image field names from table config
        image_fields = [f['name'] for f in table.fields if f.get('type') == 'image']
        for f in table.fields:
            if f['name'].upper() == 'PHOTO' and f['name'] not in image_fields:
                image_fields.append(f['name'])
        
        if not image_fields:
            return JsonResponse({'success': False, 'message': 'No image fields configured in this table!'}, status=400)
        
        # Read ZIP file and extract image data - CASE SENSITIVE keys
        zip_file = request.FILES['zip_file']
        zip_photos = {}  # Map: filename_without_ext (CASE SENSITIVE) -> {bytes, ext, original_name}
        invalid_images = 0
        
        try:
            zip_content = zip_file.read()
            with zipfile.ZipFile(BytesIO(zip_content), 'r') as zf:
                for zip_info in zf.infolist():
                    if zip_info.is_dir():
                        continue
                    
                    file_in_zip = zip_info.filename
                    base_name = os.path.basename(file_in_zip)
                    name_without_ext = os.path.splitext(base_name)[0]  # CASE SENSITIVE - no .upper()
                    ext = os.path.splitext(base_name)[1].lower()
                    
                    # Only process image files
                    if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
                        try:
                            image_bytes = zf.read(zip_info.filename)
                            
                            # Validate image before storing
                            is_valid, error_msg = validate_image_bytes(image_bytes)
                            if is_valid:
                                zip_photos[name_without_ext] = {
                                    'bytes': image_bytes,
                                    'ext': ext,
                                    'original_name': base_name
                                }
                            else:
                                invalid_images += 1
                                print(f"Invalid image skipped in reupload ZIP: {base_name} - {error_msg}")
                        except Exception as img_read_err:
                            invalid_images += 1
                            print(f"Error reading image from reupload ZIP: {base_name} - {img_read_err}")
                            continue
        except Exception as zip_error:
            return JsonResponse({'success': False, 'message': f'Error reading ZIP file: {str(zip_error)}'}, status=400)
        
        if not zip_photos:
            return JsonResponse({'success': False, 'message': 'No images found in the ZIP file!'}, status=400)
        
        # Process cards and match images using multiple matching strategies
        # Strategy 1: Match by original PHOTO column value (e.g., "1", "john", "student_001")
        # Strategy 2: Match by current stored filename (e.g., "20260203145230_001")
        images_matched = 0
        images_not_matched = 0
        cards_updated = 0
        batch_counter = 0  # Counter for this reupload batch
        
        for card in cards_to_process:
            field_data = card.field_data or {}
            card_updated = False
            
            # Collect all possible matching keys for this card
            matching_keys = []
            existing_image_path = None
            pending_reference = None  # Store the PENDING: reference if found
            
            for img_field in image_fields:
                current_value = field_data.get(img_field, '')
                
                # Check for PENDING: prefix (image reference waiting for upload)
                if current_value and current_value.startswith('PENDING:'):
                    pending_reference = current_value[8:]  # Extract the reference after "PENDING:"
                    existing_image_path = 'PENDING'
                    # Add the pending reference as a matching key
                    if pending_reference:
                        # Remove extension if present for matching
                        ref_no_ext = os.path.splitext(pending_reference)[0] if '.' in pending_reference else pending_reference
                        if ref_no_ext and ref_no_ext not in matching_keys:
                            matching_keys.append(ref_no_ext)
                        # Also add with extension
                        if pending_reference and pending_reference not in matching_keys:
                            matching_keys.append(pending_reference)
                elif current_value and current_value not in ['NOT_FOUND', '']:
                    existing_image_path = current_value
                    # Strategy 2: Add current stored filename (without extension)
                    stored_filename = os.path.splitext(os.path.basename(current_value))[0]
                    if stored_filename and stored_filename not in matching_keys:
                        matching_keys.append(stored_filename)
                elif current_value == 'NOT_FOUND':
                    # Legacy: Card has a PHOTO reference but image wasn't found during upload
                    existing_image_path = 'NOT_FOUND'
            
            # Strategy 3: Try matching by any text field that looks like a photo reference
            # (useful when PHOTO column had values like "1", "2", "john", etc.)
            for field in table.fields:
                if field['name'].upper() == 'PHOTO' or field.get('type') == 'image':
                    photo_ref = field_data.get(field['name'], '')
                    # Check for PENDING: prefix
                    if photo_ref and photo_ref.startswith('PENDING:'):
                        ref_value = photo_ref[8:]
                        ref_no_ext = os.path.splitext(ref_value)[0] if '.' in ref_value else ref_value
                        if ref_no_ext and ref_no_ext not in matching_keys:
                            matching_keys.append(ref_no_ext)
                    elif photo_ref and photo_ref != 'NOT_FOUND':
                        if '/' in photo_ref:
                            # It's a path - extract filename without extension
                            ref_name = os.path.splitext(os.path.basename(photo_ref))[0]
                        else:
                            # It's a direct reference (like "1" or "john")
                            ref_name = os.path.splitext(photo_ref)[0] if '.' in photo_ref else photo_ref
                        if ref_name and ref_name not in matching_keys:
                            matching_keys.append(ref_name)
            
            # Strategy 4: Match by card ID as fallback
            card_id_str = str(card.id)
            if card_id_str not in matching_keys:
                matching_keys.append(card_id_str)
            
            # Try to find a match in the ZIP photos using any of the matching keys
            matched_photo_info = None
            matched_key = None
            for key in matching_keys:
                if key in zip_photos:
                    matched_photo_info = zip_photos[key]
                    matched_key = key
                    break
            
            if matched_photo_info:
                try:
                    photo_info = matched_photo_info
                    batch_counter += 1
                    
                    # Generate filename based on whether there's an existing valid image
                    # PENDING and NOT_FOUND mean no actual image file exists yet
                    if existing_image_path and existing_image_path not in ['NOT_FOUND', 'PENDING', '']:
                        # UPDATE: Keep original 13-digit timestamp, add new 6-digit HHMMSS (20 digits total)
                        new_filename = generate_updated_image_filename(existing_image_path, photo_info['ext'])
                        
                        # Delete old image file before saving new one
                        try:
                            if default_storage.exists(existing_image_path):
                                default_storage.delete(existing_image_path)
                                print(f"Deleted old image during reupload: {existing_image_path}")
                        except Exception as del_err:
                            print(f"Warning: Could not delete old image {existing_image_path}: {del_err}")
                    else:
                        # FIRST UPLOAD (was PENDING/NOT_FOUND/empty): Generate fresh 13-digit filename
                        new_filename = generate_image_filename(batch_counter, photo_info['ext'])
                    
                    # Save to client's UUID folder
                    file_path = f"{client_image_folder}/{new_filename}"
                    
                    saved_path, renamed, success = safe_save_image(
                        default_storage, file_path, ContentFile(photo_info['bytes']),
                        fallback_name=f"reupload_fallback_{int(time.time())}{photo_info['ext']}"
                    )
                    
                    if success and saved_path:
                        # Update field_data with new path for ALL image fields
                        for img_field in image_fields:
                            # Update all image fields (whether empty, NOT_FOUND, or has existing image)
                            field_data[img_field] = saved_path
                            card_updated = True
                        
                        images_matched += 1
                    else:
                        print(f"Warning: Could not save reuploaded image for card {card.id}")
                except Exception as save_error:
                    # Log but don't break the whole process
                    print(f"Error saving reuploaded image for card {card.id}: {save_error}")
                    continue
            else:
                images_not_matched += 1
            
            if card_updated:
                card.field_data = field_data
                card.save()
                cards_updated += 1
        
        if images_matched == 0:
            invalid_msg = f" ({invalid_images} invalid images skipped)" if invalid_images > 0 else ""
            return JsonResponse({
                'success': False, 
                'message': f'No matching images found! ZIP contains {len(zip_photos)} valid images{invalid_msg} but none matched the stored photo references (case-sensitive).'
            }, status=400)
        
        invalid_msg = f" ({invalid_images} invalid images skipped)" if invalid_images > 0 else ""
        return JsonResponse({
            'success': True,
            'message': f'Successfully reuploaded {images_matched} images for {cards_updated} cards!{invalid_msg}',
            'images_matched': images_matched,
            'images_not_matched': images_not_matched,
            'invalid_images': invalid_images,
            'cards_updated': cards_updated,
            'zip_image_count': len(zip_photos)
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_download_docx(request, table_id):
    """API endpoint to download selected cards as Word document (.docx or .doc format)"""
    try:
        from docx import Document
        from docx.shared import Inches, Cm, Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        from io import BytesIO
        from datetime import datetime
        from django.http import HttpResponse
        from django.core.files.storage import default_storage
        import os
        from PIL import Image
        
        table = get_object_or_404(IDCardTable, id=table_id)
        data = json.loads(request.body)
        
        card_ids = data.get('card_ids', [])
        doc_format = data.get('format', 'docx')  # 'docx' or 'doc'
        
        if not card_ids:
            return JsonResponse({'success': False, 'message': 'No cards selected!'}, status=400)
        
        # Get selected cards in database order (first uploaded = first shown)
        cards = IDCard.objects.filter(table=table, id__in=card_ids).order_by('id')
        if not cards.exists():
            return JsonResponse({'success': False, 'message': 'No cards found!'}, status=400)
        
        # Get table fields configuration - maintain original order
        table_fields = table.fields
        
        # Build ordered_fields list - TEXT fields first, then IMAGE fields (images on right side)
        # Each field has: name, type, is_image flag
        text_fields = []
        image_fields = []
        for f in table_fields:
            field_name = f['name']
            field_type = f.get('type', 'text')
            # Check if it's an image field
            is_image = (field_type == 'image' or field_name.upper() in ['PHOTO', 'SIGNATURE', 'IMAGE', 'PIC', 'PICTURE', 'SIGN'])
            field_info = {
                'name': field_name,
                'type': field_type,
                'is_image': is_image
            }
            if is_image:
                image_fields.append(field_info)
            else:
                text_fields.append(field_info)
        
        # Combine: text fields first (left), image fields last (right)
        ordered_fields = text_fields + image_fields
        
        # Get client/institution name from the table's group
        institution_name = table.group.client.name if table.group and table.group.client else "Institution"
        
        # Create Word document
        doc = Document()
        
        # Set page orientation to landscape with 1cm margins
        from docx.enum.section import WD_ORIENT
        section = doc.sections[0]
        # Swap width and height for landscape
        new_width = section.page_height
        new_height = section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.orientation = WD_ORIENT.LANDSCAPE
        
        # Set 1cm margins on all sides
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.3)  # Minimal gap between content and footer
        
        # Header/Footer distance from edge
        section.header_distance = Cm(0.3)
        section.footer_distance = Cm(1)  # 1cm from bottom edge
        
        # Add Header
        header = section.header
        header.is_linked_to_previous = False
        
        # Get current date formatted
        from datetime import datetime
        current_date = datetime.now().strftime('%d-%m-%Y')
        
        # Create header table for 3-column layout - use full available width matching data table
        header_table = header.add_table(rows=1, cols=3, width=Cm(27.5))
        header_table.autofit = False
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_cells = header_table.rows[0].cells
        
        # Set column widths for header (left wider for institute name, center for title, right for brand)
        header_cells[0].width = Cm(9)
        header_cells[1].width = Cm(11)
        header_cells[2].width = Cm(7.5)
        
        # Left: INSTITUTE NAME: [Client Name] - bold, Arial
        left_para = header_cells[0].paragraphs[0]
        left_run = left_para.add_run(f'INSTITUTE NAME: {institution_name}')
        left_run.bold = True
        left_run.font.name = 'Arial'
        left_run.font.size = Pt(10)
        left_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Remove spacing below paragraph
        pPr_left = left_para._p.get_or_add_pPr()
        pPr_left.append(parse_xml(r'<w:spacing {} w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'.format(nsdecls('w'))))
        
        # Center: [Table Name] (Current Date) - bold, Arial
        center_para = header_cells[1].paragraphs[0]
        center_run = center_para.add_run(f'{table.name} ({current_date})')
        center_run.bold = True
        center_run.font.name = 'Arial'
        center_run.font.size = Pt(11)
        center_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Remove spacing below paragraph
        pPr_center = center_para._p.get_or_add_pPr()
        pPr_center.append(parse_xml(r'<w:spacing {} w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'.format(nsdecls('w'))))
        
        # Right: ADARSH ID CARDS - bold, Arial
        right_para = header_cells[2].paragraphs[0]
        right_run = right_para.add_run('ADARSH ID CARDS')
        right_run.bold = True
        right_run.font.name = 'Arial'
        right_run.font.size = Pt(10)
        right_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Remove spacing below paragraph
        pPr_right = right_para._p.get_or_add_pPr()
        pPr_right.append(parse_xml(r'<w:spacing {} w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'.format(nsdecls('w'))))
        
        # Remove header table borders and set minimal spacing + vertical center
        from docx.oxml.ns import qn
        for cell in header_cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                r'<w:tcBorders {0}>'
                r'<w:top w:val="nil"/>'
                r'<w:left w:val="nil"/>'
                r'<w:bottom w:val="nil"/>'
                r'<w:right w:val="nil"/>'
                r'</w:tcBorders>'.format(nsdecls('w'))
            )
            tcPr.append(tcBorders)
            # Add vertical alignment center
            vAlign = parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w')))
            tcPr.append(vAlign)
            # Zero cell margins
            tcMar = parse_xml(
                r'<w:tcMar {}>'
                r'<w:top w:w="0" w:type="dxa"/>'
                r'<w:bottom w:w="0" w:type="dxa"/>'
                r'</w:tcMar>'.format(nsdecls('w'))
            )
            tcPr.append(tcMar)
        
        # No underline - removed
        
        # Add Footer with 2 lines - minimal gap from content
        footer = section.footer
        footer.is_linked_to_previous = False
        
        from docx.oxml import OxmlElement
        
        # Footer Line 1 - Note text (7pt, Arial)
        footer_para1 = footer.add_paragraph()
        footer_run1 = footer_para1.add_run('Note: This document is computer generated. Please verify all details before printing ID cards.')
        footer_run1.font.name = 'Arial'
        footer_run1.font.size = Pt(7)
        footer_run1.font.color.rgb = RGBColor(0, 0, 0)
        footer_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr1 = footer_para1._p.get_or_add_pPr()
        pPr1.append(parse_xml(r'<w:spacing {} w:before="0" w:after="0" w:line="180" w:lineRule="exact"/>'.format(nsdecls('w'))))
        
        # Footer Line 2 - Generated date + copyright on left, Page X of Y on right (using tabs)
        footer_para2 = footer.add_paragraph()
        pPr2 = footer_para2._p.get_or_add_pPr()
        pPr2.append(parse_xml(r'<w:spacing {} w:before="0" w:after="0" w:line="180" w:lineRule="exact"/>'.format(nsdecls('w'))))
        
        # Add tab stop at right margin for right-aligned page number
        tabs = parse_xml(r'<w:tabs {}><w:tab w:val="right" w:pos="14400"/></w:tabs>'.format(nsdecls('w')))
        pPr2.append(tabs)
        
        # Left part: Generated date + copyright (7pt)
        left_run = footer_para2.add_run('Generated on: ' + datetime.now().strftime('%d-%b-%Y %I:%M %p') + ' | © Adarsh ID Cards Management System - All Rights Reserved')
        left_run.font.name = 'Arial'
        left_run.font.size = Pt(7)
        left_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Tab to move to right
        tab_run = footer_para2.add_run('\t')
        
        # Page X of Y (9pt, bold)
        page_run = footer_para2.add_run('Page ')
        page_run.font.name = 'Arial'
        page_run.font.size = Pt(9)
        page_run.font.bold = True
        page_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        page_num_run = footer_para2.add_run()
        page_num_run.font.size = Pt(9)
        page_num_run.font.bold = True
        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)
        page_num_run._r.append(fldChar3)
        
        of_run = footer_para2.add_run(' of ')
        of_run.font.name = 'Arial'
        of_run.font.size = Pt(9)
        of_run.font.bold = True
        of_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # NUMPAGES field
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = "NUMPAGES"
        fldChar5 = OxmlElement('w:fldChar')
        fldChar5.set(qn('w:fldCharType'), 'separate')
        fldChar6 = OxmlElement('w:fldChar')
        fldChar6.set(qn('w:fldCharType'), 'end')
        
        total_pages_run = footer_para2.add_run()
        total_pages_run.font.size = Pt(9)
        total_pages_run.font.bold = True
        total_pages_run._r.append(fldChar4)
        total_pages_run._r.append(instrText2)
        total_pages_run._r.append(fldChar5)
        total_pages_run._r.append(fldChar6)
        
        # Calculate number of columns: Sr No + all fields (in original order)
        num_cols = 1 + len(ordered_fields)
        
        # First pass: Calculate max content length for each column to determine widths
        column_max_lengths = {}
        column_max_lengths[0] = 5  # Sr No. column - fixed width
        
        # Calculate widths for each field in order
        for idx, field in enumerate(ordered_fields):
            field_name = field['name']
            if field['is_image']:
                # Image fields - fixed width for 2.5cm height images
                column_max_lengths[1 + idx] = 12
            else:
                # Text fields - calculate based on content
                max_len = len(field_name)  # Start with header length
                for card in cards:
                    field_data = card.field_data or {}
                    value = str(field_data.get(field_name, ''))
                    if len(value) > max_len:
                        max_len = len(value)
                # Cap max length for very long text
                column_max_lengths[1 + idx] = min(max_len, 50)
        
        # Calculate column widths based on content length
        # Available width in landscape A4 with 1cm margins: approximately 27.7cm
        total_chars = sum(column_max_lengths.values())
        available_width_cm = 27.5  # Full width for landscape with 1cm margins
        
        # Calculate column widths
        column_widths = {}
        for col_idx in range(num_cols):
            col_chars = column_max_lengths.get(col_idx, 10)
            col_width_cm = (col_chars / total_chars) * available_width_cm
            col_width_cm = max(1.5, min(col_width_cm, 8.0))
            column_widths[col_idx] = col_width_cm
        
        # Helper function to set cell margins (minimal padding)
        def set_cell_margins(cell, top=0, bottom=0, left=28, right=28):
            """Set cell margins in twips (1/20 of a point). 28 twips ≈ 1px"""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                r'<w:tcMar {}>'
                r'<w:top w:w="{}" w:type="dxa"/>'
                r'<w:bottom w:w="{}" w:type="dxa"/>'
                r'<w:left w:w="{}" w:type="dxa"/>'
                r'<w:right w:w="{}" w:type="dxa"/>'
                r'</w:tcMar>'.format(nsdecls('w'), top, bottom, left, right)
            )
            tcPr.append(tcMar)
        
        # Helper function to set cell vertical alignment
        def set_cell_vertical_alignment(cell, align='center'):
            """Set vertical alignment: 'top', 'center', 'bottom'"""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = parse_xml(r'<w:vAlign {} w:val="{}"/>'.format(nsdecls('w'), align))
            tcPr.append(vAlign)
        
        # Helper function to set paragraph spacing
        def set_paragraph_spacing(paragraph, before=0, after=0, line=240):
            """Set paragraph spacing in twips"""
            pPr = paragraph._p.get_or_add_pPr()
            spacing = parse_xml(
                r'<w:spacing {} w:before="{}" w:after="{}" w:line="{}" w:lineRule="auto"/>'.format(
                    nsdecls('w'), before, after, line
                )
            )
            pPr.append(spacing)
        
        # Helper function to set table borders to 0.5pt (half point)
        def set_table_borders_half_pt(table):
            """Set all table borders to 0.5pt (4 eighths of a point = sz 4)"""
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr {}/>' .format(nsdecls('w')))
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
            
            # Remove existing borders if any
            for child in tblPr:
                if 'tblBorders' in child.tag:
                    tblPr.remove(child)
                    break
            
            # sz=4 means 0.5pt (size is in eighths of a point)
            tblBorders = parse_xml(
                r'<w:tblBorders {}>'
                r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
                r'<w:left w:val="single" w:sz="4" w:color="000000"/>'
                r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
                r'<w:right w:val="single" w:sz="4" w:color="000000"/>'
                r'<w:insideH w:val="single" w:sz="4" w:color="000000"/>'
                r'<w:insideV w:val="single" w:sz="4" w:color="000000"/>'
                r'</w:tblBorders>'.format(nsdecls('w'))
            )
            tblPr.append(tblBorders)
        
        # Helper function to style header row - NO BACKGROUND, only bold, Arial font
        def style_header_row(header_cells):
            col_idx = 0
            
            # Sr No header
            header_cells[col_idx].text = 'Sr No.'
            header_cells[col_idx].paragraphs[0].runs[0].bold = True
            header_cells[col_idx].paragraphs[0].runs[0].font.name = 'Arial'
            header_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
            header_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            header_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_margins(header_cells[col_idx], 0, 0, 14, 14)
            set_cell_vertical_alignment(header_cells[col_idx], 'center')
            set_paragraph_spacing(header_cells[col_idx].paragraphs[0], 0, 0)
            header_cells[col_idx].width = Cm(column_widths[col_idx])
            col_idx += 1
            
            # All field headers in original order
            for field in ordered_fields:
                header_cells[col_idx].text = field['name']
                header_cells[col_idx].paragraphs[0].runs[0].bold = True
                header_cells[col_idx].paragraphs[0].runs[0].font.name = 'Arial'
                header_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
                header_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                header_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_margins(header_cells[col_idx], 0, 0, 14, 14)
                set_cell_vertical_alignment(header_cells[col_idx], 'center')
                set_paragraph_spacing(header_cells[col_idx].paragraphs[0], 0, 0)
                header_cells[col_idx].width = Cm(column_widths[col_idx])
                col_idx += 1
            
            # NO background shading - just bold text
        
        # Fixed row height for data rows - 2.5cm (same as image, no gap)
        row_height = Cm(2.5)
        ENTRIES_PER_PAGE = 7
        
        # Convert cards queryset to list
        cards_list = list(cards)
        total_cards = len(cards_list)
        
        # Remove default empty paragraph that Word creates
        if doc.paragraphs:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
        
        # Process cards in batches of 7 per page
        current_table = None
        sr_no = 1
        
        for card_idx, card in enumerate(cards_list):
            # Check if we need a new page/table
            if card_idx % ENTRIES_PER_PAGE == 0:
                if current_table is not None:
                    # Add page break after previous table (not as a separate paragraph)
                    # Use section break for continuous pages instead of page break
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    
                    # Add paragraph with page break
                    p = doc.add_paragraph()
                    # Remove spacing from this paragraph
                    pPr = p._p.get_or_add_pPr()
                    pPr.append(parse_xml(r'<w:spacing {} w:before="0" w:after="0" w:line="0" w:lineRule="auto"/>'.format(nsdecls('w'))))
                    
                    # Add page break run
                    run = p.add_run()
                    br = OxmlElement('w:br')
                    br.set(qn('w:type'), 'page')
                    run._r.append(br)
                
                # Create new table with header row
                current_table = doc.add_table(rows=1, cols=num_cols)
                current_table.style = 'Table Grid'
                current_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Set table borders to 0.5pt
                set_table_borders_half_pt(current_table)
                
                # Style header row
                style_header_row(current_table.rows[0].cells)
            
            # Add data row to current table
            field_data = card.field_data or {}
            new_row = current_table.add_row()
            row_cells = new_row.cells
            
            # Set fixed row height
            tr = new_row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(r'<w:trHeight {} w:val="{}" w:hRule="exact"/>'.format(
                nsdecls('w'), int(row_height.twips)
            ))
            trPr.append(trHeight)
            
            col_idx = 0
            
            # Sr No - minimal padding, Arial font
            row_cells[col_idx].text = str(sr_no)
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[col_idx].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            set_cell_margins(row_cells[col_idx], 0, 0, 14, 14)  # Minimal padding
            set_cell_vertical_alignment(row_cells[col_idx], 'center')
            set_paragraph_spacing(row_cells[col_idx].paragraphs[0], 0, 0)
            row_cells[col_idx].width = Cm(column_widths[col_idx])
            col_idx += 1
            
            # All fields in original order
            for field in ordered_fields:
                field_name = field['name']
                cell = row_cells[col_idx]
                cell.width = Cm(column_widths[col_idx])
                
                if field['is_image']:
                    # Image field - NO padding, image touches borders
                    img_path = field_data.get(field_name, '')
                    
                    # Set cell styling - zero padding on all sides
                    set_cell_margins(cell, 0, 0, 0, 0)  # No padding - image touches borders
                    set_cell_vertical_alignment(cell, 'center')
                    
                    if img_path and img_path != 'NOT_FOUND' and img_path.strip():
                        try:
                            if default_storage.exists(img_path):
                                with default_storage.open(img_path, 'rb') as img_file:
                                    img_data = img_file.read()
                                    
                                    # Validate image data is not empty
                                    if not img_data or len(img_data) < 100:
                                        raise ValueError("Image data is empty or too small")
                                    
                                    # Add 0.5pt black border to image
                                    try:
                                        pil_img = Image.open(BytesIO(img_data))
                                        
                                        # Verify image can be loaded
                                        pil_img.verify()
                                        # Re-open after verify (verify invalidates the image)
                                        pil_img = Image.open(BytesIO(img_data))
                                        
                                        # Convert to RGB if needed
                                        if pil_img.mode in ('RGBA', 'LA', 'P'):
                                            pil_img = pil_img.convert('RGB')
                                        
                                        # Add black border (1 pixel ≈ 0.5pt at 144 DPI)
                                        from PIL import ImageOps
                                        pil_img = ImageOps.expand(pil_img, border=1, fill='black')
                                        
                                        # Save image to BytesIO
                                        img_stream = BytesIO()
                                        pil_img.save(img_stream, format='JPEG', quality=90)
                                        img_stream.seek(0)
                                        
                                        # Target height is 2.5cm (fits in 2.7cm row with padding)
                                        target_height_cm = 2.5
                                        
                                        # Add image to cell
                                        paragraph = cell.paragraphs[0]
                                        run = paragraph.add_run()
                                        run.add_picture(img_stream, height=Cm(target_height_cm))
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        set_paragraph_spacing(paragraph, 0, 0)
                                    except Exception as img_err:
                                        print(f"Image processing error for {img_path}: {img_err}")
                                        cell.text = '[Error]'
                                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(150, 150, 150)
                                        set_paragraph_spacing(cell.paragraphs[0], 0, 0)
                            else:
                                cell.text = '[No Image]'
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cell.paragraphs[0].runs[0].font.size = Pt(8)
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(150, 150, 150)
                                set_paragraph_spacing(cell.paragraphs[0], 0, 0)
                        except Exception as e:
                            cell.text = '[Error]'
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if cell.paragraphs[0].runs:
                                cell.paragraphs[0].runs[0].font.size = Pt(8)
                            set_paragraph_spacing(cell.paragraphs[0], 0, 0)
                    else:
                        # Empty placeholder - leave empty (white background)
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_paragraph_spacing(paragraph, 0, 0)
                else:
                    # Text field - minimal padding, Arial font
                    value = field_data.get(field_name, '')
                    # Ensure uppercase for display
                    value = str(value).upper() if value else ''
                    cell.text = value
                    
                    # Apply minimal styling to text cells with Arial font
                    set_cell_margins(cell, 0, 0, 28, 28)  # 0 top/bottom, small left/right
                    set_cell_vertical_alignment(cell, 'center')
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.name = 'Arial'
                        cell.paragraphs[0].runs[0].font.size = Pt(9)
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_paragraph_spacing(cell.paragraphs[0], 0, 0)
                
                col_idx += 1
            
            sr_no += 1
        
        # Save document to buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Generate filename
        timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if doc_format == 'doc':
            # For .doc format, we still generate .docx but rename it
            # Most modern systems can open .docx, but for true .doc compatibility
            # we'd need additional conversion. For now, we provide .docx with .doc extension
            # which works in Word 2007+ with compatibility mode
            filename = f"{table.name}_{timestamp_str}.doc"
            content_type = 'application/msword'
        else:
            filename = f"{table.name}_{timestamp_str}.docx"
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        response = HttpResponse(doc_buffer.getvalue(), content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = len(doc_buffer.getvalue())
        
        return response
        
    except ImportError as e:
        return JsonResponse({
            'success': False, 
            'message': 'python-docx or Pillow library not installed. Run: pip install python-docx Pillow'
        }, status=500)
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def api_idcard_download_xlsx(request, table_id):
    """API endpoint to download selected cards as Excel file (.xlsx format) with auto-sized columns"""
    try:
        import openpyxl
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter
        from io import BytesIO
        from datetime import datetime
        from django.http import HttpResponse
        
        table = get_object_or_404(IDCardTable, id=table_id)
        
        # Get request data
        data = json.loads(request.body)
        card_ids = data.get('card_ids', [])
        
        if not card_ids:
            return JsonResponse({'success': False, 'message': 'No cards selected!'}, status=400)
        
        # Get cards
        cards = IDCard.objects.filter(table=table, id__in=card_ids).order_by('id')
        
        if not cards.exists():
            return JsonResponse({'success': False, 'message': 'No cards found!'}, status=400)
        
        # Get table fields
        table_fields = table.fields or []
        
        # Separate text fields (exclude image fields for Excel)
        text_fields = []
        for f in table_fields:
            field_name = f['name'].upper()
            # Skip image fields
            if field_name in ['PHOTO', 'SIGNATURE', 'IMAGE', 'PIC', 'PICTURE', 'SIGN']:
                continue
            if f.get('type', 'text') == 'image':
                continue
            text_fields.append(f)
        
        # Create workbook and worksheet
        wb = Workbook()
        ws = wb.active
        ws.title = table.name[:31]  # Excel sheet names max 31 chars
        
        # Define simple styles - no colors, just clean formatting
        header_font = Font(name='Calibri', size=11, bold=True)
        header_alignment = Alignment(horizontal='center', vertical='center')
        
        data_font = Font(name='Calibri', size=10)
        data_alignment = Alignment(horizontal='left', vertical='center', wrap_text=False)
        center_alignment = Alignment(horizontal='center', vertical='center')
        
        thin_border = Border(
            left=Side(style='thin', color='CCCCCC'),
            right=Side(style='thin', color='CCCCCC'),
            top=Side(style='thin', color='CCCCCC'),
            bottom=Side(style='thin', color='CCCCCC')
        )
        
        # Track max width for each column
        column_widths = {}
        
        # Add header row (no Sr No. - only data fields)
        headers = [f['name'] for f in text_fields]
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.alignment = header_alignment
            cell.border = thin_border
            
            # Track width
            column_widths[col_idx] = len(str(header)) + 2
        
        # Add data rows
        for row_idx, card in enumerate(cards, 2):
            field_data = card.field_data or {}
            
            # Data fields (no Sr No.)
            for col_idx, field in enumerate(text_fields, 1):
                value = field_data.get(field['name'], '')
                # Ensure uppercase for display
                value_str = str(value).upper() if value else ''
                
                cell = ws.cell(row=row_idx, column=col_idx, value=value_str)
                cell.font = data_font
                cell.alignment = data_alignment
                cell.border = thin_border
                
                # Track max width (cap at 50 to prevent overly wide columns)
                current_width = len(value_str) + 2
                current_width = min(current_width, 50)
                column_widths[col_idx] = max(column_widths.get(col_idx, 8), current_width)
        
        # Apply column widths
        for col_idx, width in column_widths.items():
            col_letter = get_column_letter(col_idx)
            # Minimum width of 8, add some padding
            final_width = max(8, width * 1.1)
            ws.column_dimensions[col_letter].width = final_width
        
        # Set row height for header
        ws.row_dimensions[1].height = 25
        
        # Freeze header row
        ws.freeze_panes = 'A2'
        
        # Save to buffer
        xlsx_buffer = BytesIO()
        wb.save(xlsx_buffer)
        xlsx_buffer.seek(0)
        
        # Generate filename
        timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{table.name}_{timestamp_str}.xlsx"
        
        response = HttpResponse(
            xlsx_buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = len(xlsx_buffer.getvalue())
        
        return response
        
    except ImportError as e:
        return JsonResponse({
            'success': False, 
            'message': 'openpyxl library not installed. Run: pip install openpyxl'
        }, status=500)
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON data!'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)
